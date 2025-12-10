"""
Document Generator for Andalusia Travel App
Generates beautiful Word documents with travel itineraries
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml.shared import OxmlElement
import io
import os
from urllib.parse import quote_plus
from datetime import datetime, timedelta
from youtube_helper import add_youtube_section_to_doc, get_video_for_city

# ============================================================================
# PATH CONFIGURATION - PORTABLE (works on any computer/cloud deployment)
# ============================================================================

# Get the directory where this file is located
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# Define paths relative to base directory
DATA_DIR = os.path.join(BASE_DIR, 'data')
PHOTOS_DIR = os.path.join(DATA_DIR, 'photos')
EVENT_PHOTOS_DIR = os.path.join(BASE_DIR, 'event_photos')

# Check if photos directory exists (for debugging deployment issues)
PHOTOS_AVAILABLE = os.path.exists(PHOTOS_DIR) and os.path.isdir(PHOTOS_DIR)
if PHOTOS_AVAILABLE:
    photo_count = len([f for f in os.listdir(PHOTOS_DIR) if f.endswith(('.jpg', '.jpeg', '.png'))])
    # print(f"✅ Photos directory found: {PHOTOS_DIR} ({photo_count} images)")
else:
    photo_count = 0
    # print(f"⚠️ Photos directory NOT found: {PHOTOS_DIR} - PDFs will not have POI images")


# ============================================================================
# AFFILIATE CONFIGURATION
# ============================================================================

# Booking.com affiliate ID (sign up at: https://partner.booking.com)
# Replace with your actual affiliate ID to earn commissions!
BOOKING_AFFILIATE_ID = "YOUR_AFFILIATE_ID"  # ⚠️ UPDATE THIS!

def get_hotel_booking_link(city_name, hotel_name=None, checkin_date=None, checkout_date=None):
    """
    Generate Booking.com affiliate link for hotel search
    
    Args:
        city_name: City to search in
        hotel_name: Optional specific hotel name
        checkin_date: Check-in date (datetime or YYYY-MM-DD string)
        checkout_date: Check-out date (datetime or YYYY-MM-DD string)
    
    Returns:
        Affiliate link URL for Booking.com
    """
    base_url = "https://www.booking.com/searchresults.html"
    
    # Build search string: hotel name + city for best results
    if hotel_name:
        search_string = f"{hotel_name} {city_name}"
    else:
        search_string = city_name
    
    # URL encode search string
    search_encoded = quote_plus(search_string)
    
    # Build parameters
    params = f"ss={search_encoded}"
    
    # Add affiliate ID if configured
    if BOOKING_AFFILIATE_ID and BOOKING_AFFILIATE_ID != "YOUR_AFFILIATE_ID":
        params += f"&aid={BOOKING_AFFILIATE_ID}"
    
    # Add dates if provided
    if checkin_date:
        if isinstance(checkin_date, datetime):
            checkin_str = checkin_date.strftime('%Y-%m-%d')
        else:
            checkin_str = str(checkin_date)
        params += f"&checkin={checkin_str}"
    
    if checkout_date:
        if isinstance(checkout_date, datetime):
            checkout_str = checkout_date.strftime('%Y-%m-%d')
        else:
            checkout_str = str(checkout_date)
        params += f"&checkout={checkout_str}"
    
    return f"{base_url}?{params}"


# ============================================================================
# HELPER FUNCTIONS
# ============================================================================


def normalize_city_name(city_name):
    """Normalize city name for matching"""
    if not city_name:
        return ""
    import unicodedata
    city_name = str(city_name)
    nfd = unicodedata.normalize('NFD', city_name)
    without_accents = ''.join(char for char in nfd if unicodedata.category(char) != 'Mn')
    return without_accents.lower().strip()


def generate_daily_map_url(previous_city, current_city, attractions, restaurants, is_circular=False, return_to_city=None):
    """
    Generate Google Maps directions URL with all POIs for the day
    
    Args:
        previous_city: Name of previous city (str or None)
        current_city: Name of current city (str)
        attractions: List of attraction dicts with 'name' and 'coordinates'
        restaurants: List of restaurant dicts (lunch, dinner)
        is_circular: Boolean indicating if this is a circular trip
        return_to_city: City to return to (for circular trips on last day)
    
    Returns:
        Google Maps URL string
    """
    from urllib.parse import quote_plus
    
    waypoints = []
    
    # Start from previous city if this is a driving day
    if previous_city and previous_city != current_city:
        waypoints.append(previous_city)
    
    # Add all attractions
    for attr in attractions:
        name = attr.get('name', '')
        coords = attr.get('coordinates', {})
        lat = coords.get('latitude') or coords.get('lat')
        lon = coords.get('longitude') or coords.get('lon') or coords.get('lng')
        
        if lat and lon:
            waypoints.append(f"{lat},{lon}")
        elif name:
            waypoints.append(name)
    
    # Add restaurants
    for restaurant in restaurants:
        if not restaurant:
            continue
        
        address = restaurant.get('address', '')
        
        if address:
            waypoints.append(address)
        else:
            # Fallback: construct from name and city
            name = restaurant.get('name', '')
            city = restaurant.get('city', '')
            if name and city:
                waypoints.append(f"{name}, {city}, Spain")
            elif name:
                waypoints.append(name)
    
    # ✅ FIX: For circular trips, add return to start city
    if is_circular and return_to_city:
        waypoints.append(return_to_city)
    
    if not waypoints:
        return None
    
    # Build Google Maps URL
    if len(waypoints) == 1:
        # Single destination
        return f"https://www.google.com/maps/dir/?api=1&destination={quote_plus(str(waypoints[0]))}"
    else:
        # Multiple waypoints
        origin = quote_plus(str(waypoints[0]))
        destination = quote_plus(str(waypoints[-1]))
        
        if len(waypoints) > 2:
            # Add intermediate waypoints
            middle_points = waypoints[1:-1]
            waypoints_param = "|".join([quote_plus(str(wp)) for wp in middle_points])
            return f"https://www.google.com/maps/dir/?api=1&origin={origin}&destination={destination}&waypoints={waypoints_param}"
        else:
            return f"https://www.google.com/maps/dir/?api=1&origin={origin}&destination={destination}"


def normalize_city_name(city_name):
    """Normalize city name for matching"""
    if not city_name:
        return ""
    import unicodedata
    city_name = str(city_name)
    nfd = unicodedata.normalize('NFD', city_name)
    without_accents = ''.join(char for char in nfd if unicodedata.category(char) != 'Mn')
    return without_accents.lower().strip()


def add_hyperlink(paragraph, url, text):
    """
    Add a working hyperlink to a Word document paragraph
    
    Args:
        paragraph: The paragraph to add the hyperlink to
        url: The URL string
        text: The display text
    
    Returns:
        The hyperlink element
    """
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import RGBColor, Pt
    
    # Get the paragraph's parent part
    part = paragraph.part
    
    # Create a relationship to the URL
    r_id = part.relate_to(
        url, 
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', 
        is_external=True
    )
    
    # Create the w:hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    # Create a new run element
    new_run = OxmlElement('w:r')
    
    # Create run properties
    rPr = OxmlElement('w:rPr')
    
    # Add color
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '2980B9')  # Blue color
    rPr.append(color)
    
    # Add underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    
    # Add the text
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    
    hyperlink.append(new_run)
    
    # Add to paragraph
    paragraph._p.append(hyperlink)
    
    return hyperlink

def get_city_prefix(city_norm):
    """Get city description"""
    descriptions = {
        'malaga': 'Coastal gem with museums, beaches, and vibrant culture',
        'granada': 'Home to the magnificent Alhambra palace',
        'cordoba': 'Historic city famous for the stunning Mezquita',
        'seville': 'Andalusia\'s capital, heart of flamenco and tapas',
        'cadiz': 'Ancient port city with beautiful beaches',
        'ronda': 'Dramatic clifftop town with iconic bridge',
        'jerez': 'Home of sherry wine and Andalusian horses',
        'tarifa': 'Southernmost point of Europe, windsurfing paradise'
    }
    return descriptions.get(city_norm, '')


def get_city_tips(city_norm):
    """Get city-specific tips"""
    tips = {
        'granada': [
            'Book Alhambra tickets 2-3 months in advance!',
            'Free tapas with every drink - bar hop in Albaicín',
            'Visit Mirador San Nicolás at sunset for stunning views'
        ],
        'seville': [
            'Alcázar is less crowded early morning',
            'Best tapas in Triana neighborhood',
            'Flamenco shows in Barrio Santa Cruz'
        ],
        'cordoba': [
            'Visit Mezquita early (8:30am) to avoid crowds',
            'Wander the flower-filled patios in spring',
            'Cross the Roman Bridge at sunset'
        ]
    }
    return tips.get(city_norm, [])


def get_poi_tip(poi_name):
    """Get POI-specific tip"""
    tips = {
        'alhambra': 'Book tickets online months in advance - they sell out!',
        'mezquita': 'Visit during morning prayer time (free entry 8:30-9:30am)',
        'alcazar': 'Book early morning slot to avoid crowds',
        'cathedral': 'Climb the tower for panoramic views'
    }
    
    name_lower = poi_name.lower()
    for key, tip in tips.items():
        if key in name_lower:
            return tip
    return None


def get_poi_description_fallback(name, category):
    """Generate fallback description if missing"""
    if not name:
        return "Interesting local attraction worth visiting"
    
    if category:
        return f"Notable {category.lower()} attraction in the area"
    
    return "Popular local point of interest"


def build_word_doc(itinerary, hop_kms, maps_link, ordered_cities, days, prefs, parsed_requests, is_car_mode=False):
    """
    Build BEAUTIFUL travel magazine-style Word document
    ... (rest of your function code)
    """
    # ... (continue with the rest of the beautiful document code)

def get_city_prefix(city_norm):
    """Get descriptive prefix for each city"""
    prefixes = {
        'malaga': '🌊 Málaga is the vibrant gateway to the Costa del Sol, birthplace of Picasso, and a perfect blend of beaches, culture, and tapas bars.',
        'sevilla': '💃 Sevilla is the heart of Andalusia, famous for flamenco, the stunning Alcázar palace, and the world\'s largest Gothic cathedral.',
        'seville': '💃 Seville is the heart of Andalusia, famous for flamenco, the stunning Alcázar palace, and the world\'s largest Gothic cathedral.',
        'granada': '🏰 Granada is home to the breathtaking Alhambra palace, nestled at the foot of the Sierra Nevada mountains.',
        'cordoba': '🕌 Córdoba boasts the magnificent Mezquita, a stunning mosque-cathedral that showcases the city\'s Moorish heritage.',
        'cadiz': '🌅 Cádiz is one of Europe\'s oldest cities, surrounded by the Atlantic Ocean with beautiful beaches and historic old town.',
        'ronda': '🌉 Ronda sits dramatically atop a gorge with the iconic Puente Nuevo bridge connecting the old and new towns.',
        'marbella': '⛱️ Marbella is a glamorous resort town on the Costa del Sol, known for luxury yachts, beaches, and upscale dining.',
        'nerja': '🏖️ Nerja is a charming coastal town famous for the Balcón de Europa viewpoint and spectacular caves.',
        'almeria': '🏜️ Almería features unique desert landscapes, historic Alcazaba fortress, and pristine Mediterranean beaches.',
        'jerez': '🍷 Jerez de la Frontera is the home of sherry wine, flamenco culture, and the Royal Andalusian School of Equestrian Art.',
        'tarifa': '🌊 Tarifa is the southernmost point of Europe, famous for windsurfing, kitesurfing, and views of Africa across the strait.',
        'gibraltar': '🗿 Gibraltar is a British territory with the famous Rock, Barbary macaques, and stunning views of two continents.',
        'antequera': '🏛️ Antequera is known for its impressive dolmens, historic churches, and the stunning El Torcal natural park.'
    }
    return prefixes.get(city_norm, f'{city_norm.title()} is a beautiful Andalusian city worth exploring.')


def get_city_tips(city_norm):
    """Get specific travel tips for each city"""
    tips = {
        'malaga': [
            'Visit the Alcazaba fortress early morning to avoid crowds and heat',
            'The Picasso Museum offers free entry in the last 2 hours on Sundays',
            'Walk along the Muelle Uno waterfront for dining and sea views',
            'Try espeto de sardinas (grilled sardines) at beach chiringuitos'
        ],
        'sevilla': [
            'Book Alcázar tickets online in advance - they sell out days ahead',
            'Visit the Cathedral early morning or late afternoon to avoid tour groups',
            'Explore Triana neighborhood across the river for authentic flamenco bars',
            'Free entry to Cathedral on Mondays for residents (show ID)'
        ],
        'seville': [
            'Book Alcázar tickets online in advance - they sell out days ahead',
            'Visit the Cathedral early morning or late afternoon to avoid tour groups',
            'Explore Triana neighborhood across the river for authentic flamenco bars',
            'Free entry to Cathedral on Mondays for residents (show ID)'
        ],
        'granada': [
            'Alhambra tickets must be booked weeks in advance - book NOW!',
            'Many bars in Albaicín offer free tapas with drinks',
            'Watch sunset from Mirador de San Nicolás for views of Alhambra',
            'Visit Alhambra in the afternoon - morning slots sell out first'
        ],
        'cordoba': [
            'Visit Mezquita first thing in the morning (8:30am) to avoid crowds',
            'Explore the Jewish Quarter (Judería) for charming patios and shops',
            'Free entry to Mezquita during morning mass hours (Mon-Sat 8:30-9:30am)',
            'Best time to visit: Spring for the Patio Festival (May)'
        ],
        'cadiz': [
            'Walk the city walls at sunset for stunning Atlantic views',
            'Visit La Caleta beach - small but charming city beach',
            'Try pescaíto frito (fried fish) at the Central Market area',
            'Explore the lively Genovés Park near the old town'
        ],
        'ronda': [
            'Visit Puente Nuevo bridge early morning for photos without crowds',
            'Walk down to the bottom of the gorge for unique bridge perspectives',
            'Try rabo de toro (oxtail stew) - a local specialty',
            'The bullring offers interesting museum tours about bullfighting history'
        ],
        'almeria': [
            'Visit the Alcazaba fortress for panoramic city and sea views',
            'Explore Cabo de Gata natural park for pristine beaches',
            'The desert landscapes were used in many Western films',
            'Try gurullos (traditional pasta dish with rabbit or seafood)'
        ],
        'jerez': [
            'Book a bodega (sherry winery) tour in advance',
            'Visit the Royal Andalusian School of Equestrian Art for horse shows',
            'Explore the Flamenco Cultural Center to learn about the dance origins',
            'Try fino or manzanilla sherry paired with local tapas'
        ],
        'tarifa': [
            'Book wind/kitesurfing lessons in advance during peak season',
            'Visit Bolonia beach for Roman ruins and pristine sand dunes',
            'Take a day trip to Tangier, Morocco (ferry departures daily)',
            'Best wind conditions: April-October for water sports'
        ],
        'nerja': [
            'Visit the Nerja Caves - spectacular stalactites and stalagmites',
            'Walk the Balcón de Europa at sunset for stunning coastal views',
            'Explore hidden beaches like Playa de Maro',
            'Try local sweet wine from the Frigiliana mountains'
        ]
    }
    return tips.get(city_norm, [])


def get_poi_tip(poi_name):
    """Get specific tips for popular POIs"""
    if not poi_name:
        return None
        
    poi_name_lower = poi_name.lower()
    
    if 'alhambra' in poi_name_lower:
        return 'Book tickets 2-3 months in advance! Morning slots sell out first. Wear comfortable shoes - lots of walking.'
    elif 'mezquita' in poi_name_lower or 'mosque' in poi_name_lower:
        return 'Visit at 8:30am for free entry during morning mass. Stunning architecture best seen in morning light.'
    elif 'alcazar' in poi_name_lower or 'alcázar' in poi_name_lower:
        return 'Book tickets online to skip long queues. Allow 2-3 hours to explore the palace and gardens thoroughly.'
    elif 'cathedral' in poi_name_lower or 'catedral' in poi_name_lower:
        return 'Climb the bell tower (Giralda in Seville) for amazing city views. Modest dress required (covered shoulders/knees).'
    elif 'picasso' in poi_name_lower:
        return 'Free entry last 2 hours on Sundays. Allow 1.5-2 hours for the full collection.'
    elif 'alcazaba' in poi_name_lower:
        return 'Visit early morning to avoid heat. Great views from the top - bring water and sun protection.'
    elif 'plaza' in poi_name_lower or 'square' in poi_name_lower:
        return 'Best visited during golden hour (sunset) for photos. Enjoy a coffee at a terrace café to soak in the atmosphere.'
    elif 'mirador' in poi_name_lower or 'viewpoint' in poi_name_lower:
        return 'Visit at sunset for magical views and photo opportunities. Can get crowded - arrive 30 minutes early.'
    elif 'beach' in poi_name_lower or 'playa' in poi_name_lower:
        return 'Pack sunscreen, water, and arrive early for best spots. Beach restaurants (chiringuitos) serve fresh seafood.'
    elif 'market' in poi_name_lower or 'mercado' in poi_name_lower:
        return 'Visit in the morning when produce is freshest. Great place to sample local foods and buy souvenirs.'
    elif 'garden' in poi_name_lower or 'jardin' in poi_name_lower:
        return 'Best visited in spring for flowers or early morning for peaceful atmosphere. Bring camera!'
    elif 'museum' in poi_name_lower or 'museo' in poi_name_lower:
        return 'Check for free entry days. Audio guides often available. Photography rules vary - check before snapping.'
    else:
        return None


def get_poi_description_fallback(poi_name, category):
    """
    Generate fallback descriptions for POIs without descriptions
    Based on category and name
    """
    if not category:
        category = "attraction"
    
    cat_lower = category.lower().strip()
    
    # Category-based templates
    category_descriptions = {
        'museum': f"A museum showcasing art, culture, and history. {poi_name} offers interesting exhibitions and collections worth exploring.",
        'museums': f"A museum showcasing art, culture, and history. {poi_name} offers interesting exhibitions and collections worth exploring.",
        'art': f"An art gallery featuring works from various artists and periods. {poi_name} is a must-visit for art enthusiasts.",
        'history': f"A historic site that tells the story of the region's past. {poi_name} provides fascinating insights into local heritage.",
        'architecture': f"An architectural landmark showcasing beautiful design and construction. {poi_name} is a stunning example of the region's architectural heritage.",
        'parks': f"A green space perfect for relaxation and outdoor activities. {poi_name} offers a peaceful escape from the city bustle.",
        'nature': f"A natural attraction featuring beautiful landscapes and scenery. {poi_name} is ideal for nature lovers and photographers.",
        'gardens': f"Beautiful gardens featuring diverse plants and landscaping. {poi_name} is perfect for a leisurely stroll.",
        'beaches': f"A coastal area with sand and sea. {poi_name} is great for swimming, sunbathing, and water activities.",
        'viewpoints': f"A scenic viewpoint offering panoramic vistas. {poi_name} provides stunning photo opportunities, especially at sunset.",
        'markets': f"A local market where you can find fresh produce, crafts, and local specialties. {poi_name} offers an authentic taste of local life.",
        'religious': f"A religious building of cultural and historical significance. {poi_name} features beautiful architecture and spiritual atmosphere.",
        'castles': f"A historic fortress showcasing medieval architecture and military history. {poi_name} offers great views and fascinating stories.",
        'palaces': f"A grand palace featuring opulent rooms and beautiful gardens. {poi_name} showcases the luxury and artistry of past eras.",
        'neighborhoods': f"A charming neighborhood with local character and atmosphere. {poi_name} is perfect for exploring on foot and discovering hidden gems.",
        'food & tapas': f"A culinary destination known for local food and flavors. {poi_name} is ideal for tasting authentic Andalusian cuisine.",
        'wine & bodegas': f"A winery or bodega offering wine tasting and tours. {poi_name} showcases the region's winemaking traditions.",
        'music & flamenco': f"A venue celebrating music and dance culture. {poi_name} offers authentic performances and cultural experiences.",
    }
    
    # Try to get category-specific description
    description = category_descriptions.get(cat_lower)
    
    # If no match, create generic description
    if not description:
        description = f"A notable {category} attraction in the area. {poi_name} is worth visiting to experience local culture and sights."
    
    return description


def build_word_doc(itinerary, hop_kms, maps_link, ordered_cities, days, prefs, parsed_requests, is_car_mode=False, result=None):
    """
    Build BEAUTIFUL travel magazine-style Word document
    
    Features:
    - Colorful headers with emojis
    - Travel-themed styling
    - Professional layout
    - Inspiring quotes
    - Beautiful formatting
    - POI photos (from local files)
    - Hub mode hotels (if result provided)
    """
    # ✅ SAFEGUARD: Ensure parsed_requests is a dict (not bool or None)
    if not isinstance(parsed_requests, dict):
        # print(f"⚠️ parsed_requests was {type(parsed_requests)}, converting to empty dict")
        parsed_requests = {}
    
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import io
    import os  # For checking local photo files
    
    # ✅ IMPROVED: Get start_date from multiple sources
    start_date = None
    try:
        # First try: get from result parameter
        if result and result.get('start_date'):
            start_date = result['start_date']
        
        # Second try: get from session state
        if start_date is None:
            import streamlit as st  # type: ignore
            start_date = st.session_state.get("current_trip_start_date")
        
        # Third try: check itinerary days for date
        if start_date is None and itinerary:
            first_day = itinerary[0] if itinerary else {}
            if first_day.get('date'):
                start_date = first_day['date']
                
    except Exception as e:
        start_date = None
    # ⚠️ DISABLED: Semantic merge was too aggressive, removing valid POIs

    # ⚠️ DISABLED: Semantic merge was too aggressive, removing valid POIs
    # The itinerary generator already handles deduplication properly
    # Applying it again here causes attractions to disappear from the Word doc
    # 
    # Original code (now disabled):
    # from semantic_merge import merge_city_pois
    # for day in itinerary:
    #     for city_stop in day.get("cities", []):
    #         attractions = city_stop.get("attractions", [])
    #         if attractions:
    #             city_name = city_stop.get("city") or day.get("city", "")
    #             city_stop["attractions"] = merge_city_pois(attractions, city_name)
    
    doc = Document()
    
    # ========================================================================
    # 🎨 STUNNING COVER PAGE
    # ========================================================================
    
    # Add some space
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Main title - HUGE and centered
    title = doc.add_heading('', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('✈️ YOUR ANDALUSIA\nROAD TRIP ADVENTURE ✈️')
    title_run.font.size = Pt(36)
    title_run.font.color.rgb = RGBColor(41, 128, 185)  # Beautiful blue
    title_run.bold = True
    
    # Subtitle with route
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    route_run = subtitle.add_run(f'🚗 {ordered_cities[0]} → {ordered_cities[-1]} 🚗')
    route_run.font.size = Pt(20)
    route_run.font.color.rgb = RGBColor(231, 76, 60)  # Vibrant red
    route_run.bold = True
    
        # Trip details box
    doc.add_paragraph()
    details = doc.add_paragraph()
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    details_text = f'📅 {days} Days  •  🏨 {len(ordered_cities)} Cities  •  {prefs.get("budget", "mid-range").title()} Budget'
    if start_date:
        try:
            from datetime import timedelta
            end_date = start_date + timedelta(days=days - 1)
            date_span = f'  •  {start_date.strftime("%d-%b-%Y")} → {end_date.strftime("%d-%b-%Y")}'
            details_text += date_span
        except Exception:
            pass
    details_run = details.add_run(details_text)
    details_run.font.size = Pt(14)
    details_run.font.color.rgb = RGBColor(52, 73, 94)

    
    # Inspiring travel quote
    doc.add_paragraph()
    doc.add_paragraph()
    quote = doc.add_paragraph()
    quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
    quote_run = quote.add_run('"The world is a book, and those who do not travel read only one page."\n– Saint Augustine')
    quote_run.font.size = Pt(12)
    quote_run.italic = True
    quote_run.font.color.rgb = RGBColor(127, 140, 141)
    
    doc.add_page_break()
    
    # ========================================================================
    # 🗺️ ROUTE OVERVIEW - Beautiful colored section
    # ========================================================================
    
    # Section header with emoji and color
    route_header = doc.add_heading('', 1)
    route_run = route_header.add_run('🗺️  YOUR ROUTE AT A GLANCE')
    route_run.font.size = Pt(24)
    route_run.font.color.rgb = RGBColor(41, 128, 185)
    route_run.bold = True
    
    # Add decorative line
    separator = doc.add_paragraph('─' * 50)
    separator_run = separator.runs[0]
    separator_run.font.color.rgb = RGBColor(189, 195, 199)
    
    # Route with arrows
    route_para = doc.add_paragraph()
    route_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    route_text = '  →  '.join(ordered_cities)
    route_text_run = route_para.add_run(f'🎯  {route_text}')
    route_text_run.font.size = Pt(14)
    route_text_run.font.color.rgb = RGBColor(52, 73, 94)
    route_text_run.bold = True
    
    doc.add_paragraph()
    
    # Trip statistics in colored boxes
    stats_para = doc.add_paragraph()
    stats_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    total_km = sum(km for km in hop_kms if km is not None)
    total_driving_hours = round(total_km / 85, 1)
    
    stats_text = f'''
    🚗  Total Driving: {int(total_km)}km (~{total_driving_hours} hours)
    📊  Average per Day: {round(total_km/days)}km
    ⛽  Estimated Fuel Cost: €{round(total_km * 0.08)}-{round(total_km * 0.12)}
    '''
    
    stats_run = stats_para.add_run(stats_text)
    stats_run.font.size = Pt(12)
    stats_run.font.color.rgb = RGBColor(52, 73, 94)
    
    # Google Maps link - highlighted
    doc.add_paragraph()
    map_para = doc.add_paragraph()
    map_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    map_icon = map_para.add_run('🌍  ')
    map_icon.font.size = Pt(14)
    
    # Add hyperlink
    # Google Maps link - highlighted
    doc.add_paragraph()
    map_para = doc.add_paragraph()
    map_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    map_icon = map_para.add_run('🌍  ')
    map_icon.font.size = Pt(14)
    
    # ✅ FIX: Use add_hyperlink function (same as daily links)
    add_hyperlink(map_para, maps_link, 'OPEN ROUTE IN GOOGLE MAPS')
    
    # Make the link bold and bigger
    for run in map_para.runs:
        if 'OPEN' in run.text:
            run.font.size = Pt(12)
            run.font.bold = True
    
    doc.add_paragraph()
    
    doc.add_paragraph()
    
    # ✅ Detect Star/Hub trip type EARLY (needed for hotel section)
    is_star_hub = False
    base_city = ''
    
    # Check if any day has is_day_trip flag (Star/Hub indicator)
    for day in itinerary:
        if day.get('is_day_trip'):
            is_star_hub = True
            base_city = day.get('base_city', '') or day.get('base', '')
            break
    
    # If not found, check if first city appears multiple times (also Star/Hub pattern)
    if not is_star_hub and itinerary:
        first_city = itinerary[0].get('city', '')
        city_counts = {}
        for day in itinerary:
            city = day.get('city', '')
            city_counts[city] = city_counts.get(city, 0) + 1
        
        # If first city appears > 50% of days, likely Star/Hub
        if first_city and city_counts.get(first_city, 0) > len(itinerary) / 2:
            is_star_hub = True
            base_city = first_city
    
    # ========================================================================
    # 🏨 HUB MODE: BASE ACCOMMODATION SECTION
    # ========================================================================
    if is_star_hub and itinerary:
        # Get base city from first day
        base_city = itinerary[0].get('base') or itinerary[0].get('city', '')
        
        # Get trip dates
        start_date_obj = None
        end_date_obj = None
        if start_date:
            start_date_obj = start_date
            # Calculate end date from itinerary length
            try:
                from datetime import timedelta
                end_date_obj = start_date + timedelta(days=len(itinerary) - 1)
            except:
                pass
        
        # Add accommodation header
        hotel_header = doc.add_heading('', 1)
        hotel_run = hotel_header.add_run('🏨  YOUR BASE ACCOMMODATION')
        hotel_run.font.size = Pt(24)
        hotel_run.font.color.rgb = RGBColor(46, 204, 113)
        hotel_run.bold = True
        
        separator = doc.add_paragraph('─' * 50)
        separator_run = separator.runs[0]
        separator_run.font.color.rgb = RGBColor(189, 195, 199)
        
        # Info about hub stay
        info_para = doc.add_paragraph()
        info_run = info_para.add_run(f'⭐ Stay in {base_city} for your entire trip. Return here each evening after day trips.')
        info_run.font.size = Pt(11)
        info_run.italic = True
        info_run.font.color.rgb = RGBColor(127, 140, 141)
        
        # Show date range
        if start_date_obj and end_date_obj:
            try:
                from datetime import timedelta
                checkin_str = start_date_obj.strftime('%d %b') if hasattr(start_date_obj, 'strftime') else str(start_date_obj)
                # ✅ FIX: Checkout is day AFTER end_date (end_date is last day of trip)
                checkout_date_display = end_date_obj + timedelta(days=1)
                checkout_str = checkout_date_display.strftime('%d %b %Y') if hasattr(checkout_date_display, 'strftime') else str(checkout_date_display)
                num_nights = (end_date_obj - start_date_obj).days + 1  # +1 because end_date is inclusive
                
                date_para = doc.add_paragraph()
                date_run = date_para.add_run(f'📅 {checkin_str} - {checkout_str} ({num_nights} nights)')
                date_run.font.size = Pt(12)
                date_run.bold = True
                date_run.font.color.rgb = RGBColor(52, 73, 94)
            except:
                pass
        
        doc.add_paragraph()
        
        # Get actual hotels from result
        base_hotels = result.get('base_hotels', []) if result else []
        
        if base_hotels:
            # Show actual hotel options
            hotel_rec_para = doc.add_paragraph()
            hotel_rec_run = hotel_rec_para.add_run(f'🔍 Recommended Hotels in {base_city}:')
            hotel_rec_run.font.size = Pt(12)
            hotel_rec_run.bold = True
            hotel_rec_run.font.color.rgb = RGBColor(44, 62, 80)
            
            doc.add_paragraph()
            
            # List each hotel with booking link
            for idx, hotel in enumerate(base_hotels, 1):
                hotel_name = hotel.get('name', 'Unknown Hotel')
                hotel_address = hotel.get('address', '')
                hotel_rating = hotel.get('rating', '')
                
                # Hotel name
                hotel_para = doc.add_paragraph()
                hotel_para.paragraph_format.left_indent = Inches(0.3)
                
                # Hotel number and name
                name_run = hotel_para.add_run(f'{idx}. {hotel_name}')
                name_run.font.size = Pt(11)
                name_run.bold = True
                name_run.font.color.rgb = RGBColor(52, 73, 94)
                
                # Rating
                if hotel_rating:
                    rating_run = hotel_para.add_run(f'  ⭐ {hotel_rating}')
                    rating_run.font.size = Pt(10)
                    rating_run.font.color.rgb = RGBColor(243, 156, 18)
                
                # Address
                if hotel_address:
                    addr_para = doc.add_paragraph()
                    addr_para.paragraph_format.left_indent = Inches(0.5)
                    addr_run = addr_para.add_run(f'📍 {hotel_address}')
                    addr_run.font.size = Pt(9)
                    addr_run.italic = True
                    addr_run.font.color.rgb = RGBColor(127, 140, 141)
                
                # Booking link with dates
                from urllib.parse import quote_plus
                hotel_search = quote_plus(f"{hotel_name} {base_city}")
                
                if start_date_obj and end_date_obj:
                    try:
                        from datetime import timedelta
                        checkin_date = start_date_obj.strftime('%Y-%m-%d') if hasattr(start_date_obj, 'strftime') else str(start_date_obj)
                        # ✅ FIX: Checkout is day AFTER end_date
                        checkout_date_obj = end_date_obj + timedelta(days=1)
                        checkout_date = checkout_date_obj.strftime('%Y-%m-%d') if hasattr(checkout_date_obj, 'strftime') else str(checkout_date_obj)
                        
                        booking_url = f"https://www.booking.com/searchresults.html?ss={hotel_search}&checkin={checkin_date}&checkout={checkout_date}"
                        if BOOKING_AFFILIATE_ID and BOOKING_AFFILIATE_ID != "YOUR_AFFILIATE_ID":
                            booking_url += f"&aid={BOOKING_AFFILIATE_ID}"
                    except:
                        booking_url = f"https://www.booking.com/search.html?ss={hotel_search}"
                else:
                    booking_url = f"https://www.booking.com/search.html?ss={hotel_search}"
                
                # Add clickable link
                link_para = doc.add_paragraph()
                link_para.paragraph_format.left_indent = Inches(0.5)
                add_hyperlink(link_para, booking_url, '      🔗 Book on Booking.com')
                
                # Style the link
                for run in link_para.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(41, 128, 185)
                
                doc.add_paragraph()  # Small space between hotels
        else:
            # Fallback: Generic booking link if no hotels provided
            hotel_rec_para = doc.add_paragraph()
            hotel_rec_run = hotel_rec_para.add_run(f'🔍 Find Hotels in {base_city}:')
            hotel_rec_run.font.size = Pt(12)
            hotel_rec_run.bold = True
            hotel_rec_run.font.color.rgb = RGBColor(44, 62, 80)
            
            # Generate generic booking link
            from urllib.parse import quote_plus
            city_encoded = quote_plus(base_city)
            
            if start_date_obj and end_date_obj:
                try:
                    checkin_date = start_date_obj.strftime('%Y-%m-%d') if hasattr(start_date_obj, 'strftime') else str(start_date_obj)
                    checkout_date = end_date_obj.strftime('%Y-%m-%d') if hasattr(end_date_obj, 'strftime') else str(end_date_obj)
                    
                    booking_url = f"https://www.booking.com/searchresults.html?ss={city_encoded}&checkin={checkin_date}&checkout={checkout_date}"
                    if BOOKING_AFFILIATE_ID and BOOKING_AFFILIATE_ID != "YOUR_AFFILIATE_ID":
                        booking_url += f"&aid={BOOKING_AFFILIATE_ID}"
                except:
                    booking_url = f"https://www.booking.com/searchresults.html?ss={city_encoded}"
            else:
                booking_url = f"https://www.booking.com/searchresults.html?ss={city_encoded}"
            
            # Add clickable booking link
            link_para = doc.add_paragraph()
            link_para.paragraph_format.left_indent = Inches(0.3)
            add_hyperlink(link_para, booking_url, f'   🔗 Book hotels in {base_city} on Booking.com')
            
            # Style the link
            for run in link_para.runs:
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(41, 128, 185)
                run.font.bold = True
        
        doc.add_paragraph()
    
    # Special requests highlighted box
    if parsed_requests.get('avoid_cities') or parsed_requests.get('must_see_cities') or parsed_requests.get('stay_duration'):
        special_box = doc.add_heading('', 2)
        special_run = special_box.add_run('🎯  YOUR PERSONALIZED PREFERENCES')
        special_run.font.size = Pt(16)
        special_run.font.color.rgb = RGBColor(230, 126, 34)
        
        if parsed_requests.get('must_see_cities'):
            must_see = doc.add_paragraph()
            must_see_text = must_see.add_run(f"✅  Must-See Cities: {', '.join(parsed_requests['must_see_cities'])}")
            must_see_text.font.color.rgb = RGBColor(39, 174, 96)
            must_see_text.font.size = Pt(11)
        
        if parsed_requests.get('avoid_cities'):
            avoid = doc.add_paragraph()
            avoid_text = avoid.add_run(f"🚫  Avoiding: {', '.join(parsed_requests['avoid_cities'])}")
            avoid_text.font.color.rgb = RGBColor(231, 76, 60)
            avoid_text.font.size = Pt(11)
        
        if parsed_requests.get('stay_duration'):
            for city, duration in parsed_requests['stay_duration'].items():
                stay = doc.add_paragraph()
                stay_text = stay.add_run(f"📅  {city}: {duration} night{'s' if duration > 1 else ''}")
                stay_text.font.color.rgb = RGBColor(52, 152, 219)
                stay_text.font.size = Pt(11)
    
    doc.add_page_break()
    
    # ========================================================================
    # 🚗 DRIVING SEGMENTS - Beautiful table-like format
    # ========================================================================
    
    # Show driving segments if we have distance data (most trips are car-based)
    if hop_kms and len(hop_kms) > 0 and any(km is not None for km in hop_kms):
        drive_header = doc.add_heading('', 1)
        drive_run = drive_header.add_run('🚗  DRIVING SEGMENTS')
        drive_run.font.size = Pt(24)
        drive_run.font.color.rgb = RGBColor(231, 76, 60)
        
        doc.add_paragraph('─' * 50)
        
        intro = doc.add_paragraph()
        if is_star_hub:
            intro_run = intro.add_run(f'⭐ Day trips from your base in {base_city} - all drives include return journey! Times are approximate. ☕📸')
        else:
            intro_run = intro.add_run('Your scenic drives through Andalusia - times are approximate, not including stops for photos, coffee, or impromptu adventures! ☕📸')
        intro_run.font.size = Pt(10)
        intro_run.italic = True
        intro_run.font.color.rgb = RGBColor(127, 140, 141)
        
        doc.add_paragraph()
        
        # ✅ For Star/Hub, show day trips from base (not city-to-city)
        if is_star_hub:
            for day in itinerary:
                if day.get('is_day_trip'):
                    trip_city = day.get('city')
                    km = day.get('driving_km', 0) / 2  # One-way distance
                    
                    if km <= 0:
                        continue
                    
                    hours = round(day.get('driving_hours', 0) / 2, 1)  # One-way time
                    
                    # Round trip format
                    seg_para = doc.add_paragraph()
                    
                    # Outbound
                    arrow = seg_para.add_run(f'  {base_city}  ')
                    arrow.font.size = Pt(12)
                    arrow.font.color.rgb = RGBColor(52, 73, 94)
                    arrow.bold = True
                    
                    arrow_symbol = seg_para.add_run(' ━━━━━━━━➤  ')
                    arrow_symbol.font.color.rgb = RGBColor(41, 128, 185)
                    arrow_symbol.font.size = Pt(12)
                    
                    destination = seg_para.add_run(f'{trip_city}  ')
                    destination.font.size = Pt(12)
                    destination.font.color.rgb = RGBColor(52, 73, 94)
                    destination.bold = True
                    
                    # Return arrow
                    return_arrow = seg_para.add_run(' ━━━━━━━━➤  ')
                    return_arrow.font.color.rgb = RGBColor(46, 204, 113)
                    return_arrow.font.size = Pt(12)
                    
                    return_dest = seg_para.add_run(f'{base_city}')
                    return_dest.font.size = Pt(12)
                    return_dest.font.color.rgb = RGBColor(52, 73, 94)
                    return_dest.bold = True
                    
                    # Distance and time (round trip)
                    details_para = doc.add_paragraph()
                    details_para.paragraph_format.left_indent = Inches(0.5)
                    
                    total_km = round(km * 2, 1)
                    total_hours = round(hours * 2, 1)
                    details_text = details_para.add_run(f'        📍  {total_km}km (round trip)  •  ⏱️  ~{total_hours}h total  •  ⛽ €{round(total_km * 0.10)} fuel')
                    details_text.font.size = Pt(10)
                    details_text.font.color.rgb = RGBColor(127, 140, 141)
                    
                    doc.add_paragraph()
        else:
            # Original city-to-city segments for Point-to-Point and Circular
            for i in range(len(ordered_cities) - 1):
                from_city = ordered_cities[i]
                to_city = ordered_cities[i + 1]
                km = hop_kms[i] if i < len(hop_kms) else None
                
                if km is None:
                    continue  # Skip if no distance data for this leg
                
                hours = round(km / 85, 1) if km else None
                
                # Each segment as a beautiful formatted block
                seg_para = doc.add_paragraph()
                
                # Route arrow
                arrow = seg_para.add_run(f'  {from_city}  ')
                arrow.font.size = Pt(12)
                arrow.font.color.rgb = RGBColor(52, 73, 94)
                arrow.bold = True
                
                arrow_symbol = seg_para.add_run(' ━━━━━━━━➤  ')
                arrow_symbol.font.color.rgb = RGBColor(41, 128, 185)
                arrow_symbol.font.size = Pt(12)
                
                destination = seg_para.add_run(f'{to_city}')
                destination.font.size = Pt(12)
                destination.font.color.rgb = RGBColor(52, 73, 94)
                destination.bold = True
                
                # Distance and time
                details_para = doc.add_paragraph()
                details_para.paragraph_format.left_indent = Inches(0.5)
                
                details_text = details_para.add_run(f'        📍  {km}km  •  ⏱️  ~{hours}h drive  •  ⛽ €{round(km * 0.10)} fuel')
                details_text.font.size = Pt(10)
                details_text.font.color.rgb = RGBColor(127, 140, 141)
                
                doc.add_paragraph()
        
        # Driving tips box
        tips_box = doc.add_paragraph()
        tips_box.paragraph_format.left_indent = Inches(0.3)
        tips_box.paragraph_format.right_indent = Inches(0.3)
        
        tips_title = tips_box.add_run('\n💡  PRO TIP: ')
        tips_title.font.bold = True
        tips_title.font.color.rgb = RGBColor(243, 156, 18)
        tips_title.font.size = Pt(11)
        
        tips_text = tips_box.add_run('Add 20-30% extra time for rest stops, tolls, scenic viewpoints, and those "just one more photo" moments. Highways (autopistas) are fast but have tolls. Secondary roads are slower but more scenic!')
        tips_text.font.size = Pt(10)
        tips_text.font.color.rgb = RGBColor(52, 73, 94)
        tips_text.italic = True
        
        doc.add_page_break()
    
    # ========================================================================
    # 📅 DAILY ITINERARY - Magazine-style with colors and emojis
    # ========================================================================
    
    itinerary_header = doc.add_heading('', 1)
    itinerary_run = itinerary_header.add_run('📅  YOUR DAY-BY-DAY ADVENTURE')
    itinerary_run.font.size = Pt(24)
    itinerary_run.font.color.rgb = RGBColor(155, 89, 182)
    
    doc.add_paragraph('─' * 50)
    doc.add_paragraph()
    
    visited_cities = set()
    
    for idx, day in enumerate(itinerary):
        city = day.get("city", "?")
        city_norm = normalize_city_name(city)
        is_must_see = day.get("is_must_see", False)
        driving_km = day.get("driving_km", 0)
        driving_hours = day.get("driving_hours", 0)
        
        # ═══════════════════════════════════════════════════════════════
        # CITY GUIDE (only once per city)
        # ═══════════════════════════════════════════════════════════════
        
        if city_norm not in visited_cities:
            visited_cities.add(city_norm)
            
            # Don't add page break before first city
            if len(visited_cities) > 1:
                doc.add_page_break()
            
            # Big colorful city header
            city_header = doc.add_heading('', 0)
            city_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # City name in beautiful color
            city_title = city_header.add_run(f'📍  {city.upper()}')
            city_title.font.size = Pt(32)
            city_title.font.color.rgb = RGBColor(41, 128, 185)
            city_title.bold = True
            
            # Must-see badge
            if is_must_see:
                must_see_para = doc.add_paragraph()
                must_see_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                must_see_run = must_see_para.add_run('⭐ MUST-SEE DESTINATION ⭐')
                must_see_run.font.size = Pt(14)
                must_see_run.font.color.rgb = RGBColor(243, 156, 18)
                must_see_run.bold = True
            
            # City description
            city_desc = get_city_prefix(city_norm)
            if city_desc:
                desc_para = doc.add_paragraph()
                desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                desc_run = desc_para.add_run(city_desc)
                desc_run.font.size = Pt(12)
                desc_run.italic = True
                desc_run.font.color.rgb = RGBColor(127, 140, 141)
            
            doc.add_paragraph()
            # 🎬 YOUTUBE VIDEO PREVIEW
            try:
                video_added = add_youtube_section_to_doc(doc, city, add_hyperlink)
                if video_added:
                    doc.add_paragraph()
            except Exception as e:
                print(f"⚠️ Could not add YouTube video for {city}: {e}")
            doc.add_paragraph('═' * 50)
            doc.add_paragraph()
            
            # City tips in colored box
            city_tips = get_city_tips(city_norm)
            if city_tips:
                tips_header = doc.add_heading('', 2)
                tips_run = tips_header.add_run('💡  LOCAL INSIDER TIPS')
                tips_run.font.size = Pt(16)
                tips_run.font.color.rgb = RGBColor(243, 156, 18)
                
                for tip in city_tips:
                    tip_para = doc.add_paragraph(style='List Bullet')
                    tip_run = tip_para.add_run(tip)
                    tip_run.font.size = Pt(10)
                    tip_run.font.color.rgb = RGBColor(52, 73, 94)
                
                doc.add_paragraph()
        
        # ═══════════════════════════════════════════════════════════════
        # DAY HEADER
        # ═══════════════════════════════════════════════════════════════
        
                # ═══════════════════════════════════════════════════════════════
        # DAY HEADER
        # ═══════════════════════════════════════════════════════════════
        
        day_header = doc.add_heading('', 2)

        # Day number with optional calendar date
        day_num = day.get("day", 0)
        
        # ✅ NEW: Get day X of Y in city
        day_in_city = day.get("day_in_city", 1)
        total_days_in_city = day.get("total_days_in_city", 1)
        city_day_suffix = ""
        if total_days_in_city > 1:
            city_day_suffix = f" (Day {day_in_city}/{total_days_in_city})"
        
        if start_date and day_num:
            from datetime import timedelta
            trip_date = start_date + timedelta(days=day_num - 1)
            date_str = trip_date.strftime("%a, %d-%b-%Y")  # e.g. Tue, 25-Aug-2026
            day_label = f'📆  DAY {day_num}: {date_str} – {city}{city_day_suffix}'
        else:
            day_label = f'📆  DAY {day_num}: {city}{city_day_suffix}'
        day_title = day_header.add_run(day_label)
        day_title.font.size = Pt(20)
        day_title.font.color.rgb = RGBColor(52, 152, 219)
        day_title.bold = True

        
        # Driving info box if applicable
        if driving_km > 0:
            drive_box = doc.add_paragraph()
            drive_box.paragraph_format.left_indent = Inches(0.5)
            
            drive_emoji = drive_box.add_run('🚗  ')
            drive_emoji.font.size = Pt(12)
            
            drive_text = drive_box.add_run(f'Drive: {driving_km}km (~{driving_hours}h)')
            drive_text.font.size = Pt(11)
            drive_text.font.color.rgb = RGBColor(231, 76, 60)
            drive_text.bold = True
            
            drive_tip = drive_box.add_run('  •  Leave early to beat traffic!')
            drive_tip.font.size = Pt(9)
            drive_tip.italic = True
            drive_tip.font.color.rgb = RGBColor(127, 140, 141)
        
        doc.add_paragraph()
        
        # ═══════════════════════════════════════════════════════════════
        # GOOGLE MAPS LINK - Route for the day
        # ═══════════════════════════════════════════════════════════════
        
        # Collect all stops for the day
        day_attractions = []
        for city_stop in day.get("cities", []):
            day_attractions.extend(city_stop.get("attractions", []))
        
        day_restaurants = [
            day.get("lunch_restaurant"),
            day.get("dinner_restaurant")
        ]
        day_restaurants = [r for r in day_restaurants if r]  # Remove None values
        
        # Get previous city for driving days
        prev_city = None
        if day.get("day", 1) > 1:
            # Find previous day's city
            for prev_day in itinerary:
                if prev_day.get("day") == day.get("day") - 1:
                    prev_city = prev_day.get("overnight_city") or prev_day.get("city")
                    break
        
        # ✅ FIX: Detect if this is last day of circular trip
        is_last_day = (idx == len(itinerary) - 1)
        overnight = day.get("overnight_city", city)
        is_circular_return = (is_last_day and city != overnight)
        
        # Generate map URL
        if day_attractions or day_restaurants:
            map_url = generate_daily_map_url(
                prev_city, 
                city, 
                day_attractions, 
                day_restaurants,
                is_circular=is_circular_return,
                return_to_city=overnight if is_circular_return else None
            )
            
            if map_url:
                map_para = doc.add_paragraph()
                map_para.paragraph_format.left_indent = Inches(0.3)
                
                map_icon = map_para.add_run('🗺️  ')
                map_icon.font.size = Pt(11)
                
                map_label = map_para.add_run('Today\'s Route:  ')
                map_label.font.size = Pt(10)
                map_label.font.color.rgb = RGBColor(52, 73, 94)
                
                # Add hyperlink
                add_hyperlink(map_para, map_url, 'Open in Google Maps')
                
                doc.add_paragraph()
        
        # ═══════════════════════════════════════════════════════════════
        # ATTRACTIONS - Beautiful cards
        # ═══════════════════════════════════════════════════════════════
        
        cities_list = day.get("cities", [])
        for city_stop in cities_list:
            attractions = city_stop.get("attractions", [])
            
            if attractions:
                attr_header = doc.add_heading('', 3)
                attr_run = attr_header.add_run('🎯  TODAY\'S HIGHLIGHTS')
                attr_run.font.size = Pt(14)
                attr_run.font.color.rgb = RGBColor(155, 89, 182)
                
                for idx, attr in enumerate(attractions, 1):
                    attr_name = attr.get('name', '?')
                    if attr.get('is_must_see_attraction'):
                        attr_name += ' ⭐'
                    
                    # Attraction name - bold and colorful
                    attr_para = doc.add_paragraph()
                    
                    number = attr_para.add_run(f'{idx}. ')
                    number.font.size = Pt(12)
                    number.font.color.rgb = RGBColor(52, 152, 219)
                    number.bold = True
                    
                    name_run = attr_para.add_run(attr_name)
                    name_run.bold = True
                    name_run.font.size = Pt(12)
                    name_run.font.color.rgb = RGBColor(44, 62, 80)
                    
                    # ✅ NEW: Add photo if available
                    # Try local_photo_path first, then construct from photo_references
                    photo_path = None
                    
                    # Option 1: Use local_photo_path if available
                    if attr.get('local_photo_path'):
                        relative_path = attr['local_photo_path']
                        
                        # ✅ FIX: Normalize path separators (Windows \ to Unix /)
                        relative_path = relative_path.replace('\\', '/')
                        
                        # Convert relative path to absolute
                        if not os.path.isabs(relative_path):
                            # It's relative (like "photos/ChIJ...jpg")
                            # Convert using DATA_DIR (portable!)
                            photo_path = os.path.join(DATA_DIR, relative_path)
                        else:
                            # Already absolute
                            photo_path = relative_path
                        
                    
                    # Option 2: Construct path from photo_references (SMART!)
                    elif attr.get('photo_references'):
                        # Use place_id to find the photo file
                        place_id = attr.get('place_id')
                        if place_id:
                            # Photo filename is place_id + .jpg
                            photo_filename = f"{place_id}.jpg"
                            
                            # PORTABLE PATH - Works on any system!
                            photo_path = os.path.join(PHOTOS_DIR, photo_filename)
                    
                    # ✅ FIX: Also try just the filename in PHOTOS_DIR as fallback
                    if photo_path and not os.path.exists(photo_path):
                        # Try extracting just the filename and looking in PHOTOS_DIR
                        filename_only = os.path.basename(photo_path)
                        fallback_path = os.path.join(PHOTOS_DIR, filename_only)
                        if os.path.exists(fallback_path):
                            photo_path = fallback_path
                    
                    # Add photo if we found it
                    if photo_path and os.path.exists(photo_path):
                        try:
                            
                            photo_para = doc.add_paragraph()
                            photo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            photo_para.paragraph_format.space_before = Pt(6)
                            photo_para.paragraph_format.space_after = Pt(6)
                            
                            # Add picture with width of 4.5 inches
                            run = photo_para.add_run()
                            run.add_picture(photo_path, width=Inches(4.5))
                            
                            
                        except Exception as e:
                            # If photo fails, just continue without it
                            
                            pass
                    # elif photo_path:
                      #  print(f"DEBUG: Photo path exists but file not found: {photo_path}")
                    
                    # Description
                    description = attr.get('description')
                    if not description or description.strip() == '':
                        description = get_poi_description_fallback(attr_name, attr.get('category'))
                    
                    if description:
                        desc_para = doc.add_paragraph()
                        desc_para.paragraph_format.left_indent = Inches(0.3)
                        desc_run = desc_para.add_run(description)
                        desc_run.font.size = Pt(10)
                        desc_run.font.color.rgb = RGBColor(52, 73, 94)
                    
                    # Details with icons
                    details = []
                    
                    # ✅ ALWAYS show rating
                    if attr.get('rating'):
                        details.append(f"⭐ {attr['rating']}")
                    else:
                        details.append("⭐ N/A")
                    
                    if attr.get('visit_duration_hours'):
                        details.append(f"⏱️ {attr['visit_duration_hours']}h")
                    
                    # ✅ FIX: Always show entrance fee with € symbol (not $)
                    entrance_fee = attr.get('entrance_fee', '')
                    entrance_fee_value = attr.get('entrance_fee_value')
                    
                    if entrance_fee_value is not None and entrance_fee_value > 0:
                        # Numeric value - format as euros
                        details.append(f"💶 €{entrance_fee_value}")
                    elif entrance_fee:
                        # Text value - clean up any $ symbols
                        fee_text = str(entrance_fee).replace('$', '€')
                        if not fee_text.startswith('€') and not any(x in fee_text.lower() for x in ['free', 'varies', '€']):
                            fee_text = f"€{fee_text}"
                        details.append(f"💶 {fee_text}")
                    else:
                        # No fee info - show varies
                        details.append("💶 Varies")
                    
                    if attr.get('category'):
                        details.append(f"🏷️ {attr['category'].title()}")
                    
                    if details:
                        details_para = doc.add_paragraph()
                        details_para.paragraph_format.left_indent = Inches(0.3)
                        details_run = details_para.add_run('   ' + '  •  '.join(details))
                        details_run.font.size = Pt(9)
                        details_run.font.color.rgb = RGBColor(149, 165, 166)
                    
                    # POI tip
                    poi_tip = get_poi_tip(attr_name)
                    if poi_tip:
                        tip_para = doc.add_paragraph()
                        tip_para.paragraph_format.left_indent = Inches(0.3)
                        
                        tip_icon = tip_para.add_run('💡 ')
                        tip_icon.font.size = Pt(9)
                        
                        tip_run = tip_para.add_run(poi_tip)
                        tip_run.italic = True
                        tip_run.font.size = Pt(9)
                        tip_run.font.color.rgb = RGBColor(243, 156, 18)
                    
                    doc.add_paragraph()  # Spacing
        
        # ═══════════════════════════════════════════════════════════════
        # HOTELS - Colorful recommendation boxes
        # ═══════════════════════════════════════════════════════════════
        
        hotels = day.get("hotels", [])
        overnight = day.get("overnight_city", city)
        
        # ✅ FIX: For circular trips, detect if this is the last day returning to start
        is_last_day = (idx == len(itinerary) - 1)
        is_return_to_start = (is_last_day and city != overnight)
        
        if hotels and any(h.get('name') and 'Hotels in' not in h.get('name', '') for h in hotels):
            hotel_header = doc.add_heading('', 3)
            
            # ✅ FIX: Better wording for circular trip return
            if is_return_to_start:
                hotel_run = hotel_header.add_run(f'🏨  WHERE TO STAY TONIGHT ({overnight.upper()})')
            else:
                hotel_run = hotel_header.add_run(f'🏨  WHERE TO STAY IN {overnight.upper()}')
            
            hotel_run.font.size = Pt(14)
            hotel_run.font.color.rgb = RGBColor(46, 204, 113)
            
            # ✅ FIX: Add return message for circular trips
            if is_return_to_start:
                return_para = doc.add_paragraph()
                return_para.paragraph_format.left_indent = Inches(0.3)
                return_icon = return_para.add_run('🔄  ')
                return_icon.font.size = Pt(11)
                return_text = return_para.add_run(f'Drive back to {overnight} for your final night')
                return_text.font.size = Pt(10)
                return_text.italic = True
                return_text.font.color.rgb = RGBColor(52, 152, 219)
                doc.add_paragraph()  # Spacing
            
            for hotel in hotels[:3]:
                if not hotel.get('name') or 'Hotels in' in hotel.get('name', ''):
                    continue
                
                hotel_para = doc.add_paragraph()
                hotel_para.paragraph_format.left_indent = Inches(0.3)
                
                bullet = hotel_para.add_run('🏨  ')
                bullet.font.size = Pt(12)
                
                hotel_name = hotel_para.add_run(hotel.get('name', '?'))
                hotel_name.bold = True
                hotel_name.font.size = Pt(11)
                hotel_name.font.color.rgb = RGBColor(44, 62, 80)
                
                rating = hotel.get("guest_rating") or hotel.get("star_rating")
                price = hotel.get("avg_price_per_night_couple")
                
                # NOTE: Hotel data (OpenStreetMap) doesn't include review counts
                # Only show rating without review count
                
                details = []
                if rating:
                    # Format rating appropriately based on scale
                    if rating <= 5:
                        # Star rating (1-5 scale) - show as stars
                        stars = "⭐" * int(rating)
                        details.append(stars)
                    else:
                        # Guest rating (typically 0-10 scale)
                        details.append(f"⭐ {rating}/10")
                
                if price:
                    details.append(f"€{price}/night")
                
                # Add details if any exist
                if details:
                    hotel_para.add_run(f"  •  {' • '.join(details)}")
                elif not rating and not price:
                    # If no data at all, add subtle note
                    note = hotel_para.add_run("  •  Check reviews online")
                    note.font.size = Pt(9)
                    note.italic = True
                    note.font.color.rgb = RGBColor(149, 165, 166)
                
                # ✅ NEW: Add clickable booking link
                city_name = day.get('city', '')
                
                # Calculate check-in/check-out dates
                checkin = None
                checkout = None
                
                # Get number of nights from day data
                day_in_city = day.get('day_in_city', 1)
                total_days_in_city = day.get('total_days_in_city', 1)
                
                # Only calculate booking dates for first day in city
                if day_in_city == 1:
                    # Nights = days in city (same for ALL cities)
                    nights = total_days_in_city
                    
                    # Try to get date from the day object first (most reliable)
                    day_date_obj = day.get('date_obj')  # datetime object added by trip_planner_page
                    day_date_str = day.get('date')  # String format "2026-04-10"
                    
                    if day_date_obj:
                        # Best case: we have a datetime object
                        checkin = day_date_obj
                        checkout = checkin + timedelta(days=nights)
                    elif day_date_str:
                        # Parse the date string
                        try:
                            from datetime import datetime
                            checkin = datetime.strptime(day_date_str, '%Y-%m-%d')
                            checkout = checkin + timedelta(days=nights)
                        except:
                            pass
                    
                    # Fallback: calculate from start_date if available
                    if not checkin and start_date:
                        try:
                            # Check-in is the night of this day (idx is 0-based)
                            checkin = start_date + timedelta(days=idx)
                            checkout = checkin + timedelta(days=nights)
                        except:
                            pass
                
                # Generate booking link
                booking_url = get_hotel_booking_link(
                    city_name,
                    hotel_name=hotel.get('name'),
                    checkin_date=checkin,
                    checkout_date=checkout
                )
                
                # Add booking link as clickable hyperlink
                link_para = doc.add_paragraph()
                link_para.paragraph_format.left_indent = Inches(0.6)
                link_para.paragraph_format.space_after = Pt(6)
                
                add_hyperlink(link_para, booking_url, '      🔗 Book on Booking.com')
                
                # Style the link
                for run in link_para.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(41, 128, 185)  # Blue
            
            # Parking tip
            parking = doc.add_paragraph()
            parking.paragraph_format.left_indent = Inches(0.3)
            parking_icon = parking.add_run('🅿️  ')
            parking_icon.font.size = Pt(10)
            parking_text = parking.add_run('Most hotels offer parking €10-20/night. Always ask when booking!')
            parking_text.font.size = Pt(9)
            parking_text.italic = True
            parking_text.font.color.rgb = RGBColor(127, 140, 141)
            
            doc.add_paragraph()
        
        # ═══════════════════════════════════════════════════════════════
        # RESTAURANTS - Delicious looking section
        # ═══════════════════════════════════════════════════════════════
        
        lunch = day.get("lunch_restaurant")
        dinner = day.get("dinner_restaurant")
        
        if lunch or dinner:
            food_header = doc.add_heading('', 3)
            food_run = food_header.add_run('🍽️  WHERE TO EAT TODAY')
            food_run.font.size = Pt(14)
            food_run.font.color.rgb = RGBColor(230, 126, 34)
            
            if lunch:
                lunch_para = doc.add_paragraph()
                lunch_para.paragraph_format.left_indent = Inches(0.3)
                
                lunch_icon = lunch_para.add_run('🥘  LUNCH: ')
                lunch_icon.font.bold = True
                lunch_icon.font.color.rgb = RGBColor(230, 126, 34)
                lunch_icon.font.size = Pt(11)
                
                lunch_name = lunch_para.add_run(lunch.get('name', 'Local restaurant'))
                lunch_name.font.size = Pt(11)
                lunch_name.font.color.rgb = RGBColor(44, 62, 80)
                
                # Show rating with review count if available
                if lunch.get('rating'):
                    rating_text = f"  •  ⭐ {lunch['rating']}"
                    
                    # Try to parse review count from "topic" field
                    topic = lunch.get('topic', '')
                    if topic:
                        parts = topic.split()
                        if parts and parts[0].isdigit():
                            reviews_count = int(parts[0])
                            rating_text += f" ({reviews_count} reviews)"
                    
                    lunch_para.add_run(rating_text)
                
                # ✅ FIX: Show price range if available
                price_range = lunch.get('price_range') or lunch.get('budget')
                if price_range:
                    lunch_para.add_run(f"  •  {price_range}")
                
                # ✅ FIX: Show cuisine if available
                cuisine = lunch.get('cuisine') or lunch.get('category')
                if cuisine:
                    lunch_para.add_run(f"  •  {cuisine}")
                
                # ✅ FIX: Always show address (with better formatting)
                address = lunch.get('address', '')
                if address:
                    # Clean address - prefer street address over generic city names
                    if ',' in address and 'Andalusia, Spain' in address:
                        # Show full address
                        addr_para = doc.add_paragraph()
                        addr_para.paragraph_format.left_indent = Inches(0.5)
                        addr_icon = addr_para.add_run('📍  ')
                        addr_icon.font.size = Pt(9)
                        addr_text = addr_para.add_run(address)
                        addr_text.font.size = Pt(9)
                        addr_text.font.color.rgb = RGBColor(127, 140, 141)
                    elif not any(x in address for x in ['Andalusia', 'Spain']):
                        # Just restaurant name with city - add city
                        city_name = lunch.get('city', current_city)
                        full_address = f"{address}, {city_name}"
                        addr_para = doc.add_paragraph()
                        addr_para.paragraph_format.left_indent = Inches(0.5)
                        addr_icon = addr_para.add_run('📍  ')
                        addr_icon.font.size = Pt(9)
                        addr_text = addr_para.add_run(full_address)
                        addr_text.font.size = Pt(9)
                        addr_text.font.color.rgb = RGBColor(127, 140, 141)
                    else:
                        addr_para = doc.add_paragraph()
                        addr_para.paragraph_format.left_indent = Inches(0.5)
                        addr_icon = addr_para.add_run('📍  ')
                        addr_icon.font.size = Pt(9)
                        addr_text = addr_para.add_run(address)
                        addr_text.font.size = Pt(9)
                        addr_text.font.color.rgb = RGBColor(127, 140, 141)
                
                if lunch.get('description'):
                    desc = doc.add_paragraph()
                    desc.paragraph_format.left_indent = Inches(0.5)
                    desc_run = desc.add_run(lunch['description'])
                    desc_run.font.size = Pt(9)
                    desc_run.italic = True
                    desc_run.font.color.rgb = RGBColor(127, 140, 141)
                
                doc.add_paragraph()
            
            if dinner:
                dinner_para = doc.add_paragraph()
                dinner_para.paragraph_format.left_indent = Inches(0.3)
                
                dinner_icon = dinner_para.add_run('🌙  DINNER: ')
                dinner_icon.font.bold = True
                dinner_icon.font.color.rgb = RGBColor(142, 68, 173)
                dinner_icon.font.size = Pt(11)
                
                dinner_name = dinner_para.add_run(dinner.get('name', 'Local restaurant'))
                dinner_name.font.size = Pt(11)
                dinner_name.font.color.rgb = RGBColor(44, 62, 80)
                
                # Show rating with review count if available
                if dinner.get('rating'):
                    rating_text = f"  •  ⭐ {dinner['rating']}"
                    
                    # Try to parse review count from "topic" field
                    topic = dinner.get('topic', '')
                    if topic:
                        parts = topic.split()
                        if parts and parts[0].isdigit():
                            reviews_count = int(parts[0])
                            rating_text += f" ({reviews_count} reviews)"
                    
                    dinner_para.add_run(rating_text)
                
                # ✅ FIX: Show price range if available
                price_range = dinner.get('price_range') or dinner.get('budget')
                if price_range:
                    dinner_para.add_run(f"  •  {price_range}")
                
                # ✅ FIX: Show cuisine if available
                cuisine = dinner.get('cuisine') or dinner.get('category')
                if cuisine:
                    dinner_para.add_run(f"  •  {cuisine}")
                
                # ✅ FIX: Always show address (with better formatting)
                address = dinner.get('address', '')
                if address:
                    addr_para = doc.add_paragraph()
                    addr_para.paragraph_format.left_indent = Inches(0.5)
                    addr_icon = addr_para.add_run('📍  ')
                    addr_icon.font.size = Pt(9)
                    addr_text = addr_para.add_run(address)
                    addr_text.font.size = Pt(9)
                    addr_text.font.color.rgb = RGBColor(127, 140, 141)
                
                if dinner.get('description'):
                    desc = doc.add_paragraph()
                    desc.paragraph_format.left_indent = Inches(0.5)
                    desc_run = desc.add_run(dinner['description'])
                    desc_run.font.size = Pt(9)
                    desc_run.italic = True
                    desc_run.font.color.rgb = RGBColor(127, 140, 141)
        
        # Day separator
        doc.add_paragraph()
        doc.add_paragraph('─' * 50)
        doc.add_paragraph()
        # ========================================================================
    # 🍽️ ANDALUSIAN FOOD GUIDE - Delicious section
    # ========================================================================
    
    doc.add_page_break()
    
    food_guide_header = doc.add_heading('', 1)
    food_guide_run = food_guide_header.add_run('🍽️  MUST-TRY ANDALUSIAN DISHES')
    food_guide_run.font.size = Pt(24)
    food_guide_run.font.color.rgb = RGBColor(230, 126, 34)
    
    doc.add_paragraph('═' * 50)
    
    intro = doc.add_paragraph()
    intro_run = intro.add_run('Andalusian cuisine is a delicious blend of Mediterranean and Moorish influences. Don\'t leave without trying these iconic dishes!')
    intro_run.font.size = Pt(11)
    intro_run.italic = True
    intro_run.font.color.rgb = RGBColor(127, 140, 141)
    
    doc.add_paragraph()
    
    dishes = [
        ('🥘 Gazpacho', 'Cold tomato soup, perfect for hot days - originated right here in Andalusia'),
        ('🍲 Salmorejo', 'Thick, creamy cold soup from Córdoba, topped with egg and jamón'),
        ('🥓 Jamón Ibérico', 'Premium cured ham from acorn-fed pigs - the best in Spain!'),
        ('🐟 Pescaíto Frito', 'Fried fish platter, a coastal specialty in Málaga and Cádiz'),
        ('🍖 Rabo de Toro', 'Oxtail stew, traditional in Córdoba (especially after bullfights)'),
        ('🥚 Tortilla Española', 'Classic Spanish potato omelet - simple but delicious'),
        ('🐟 Espeto de Sardinas', 'Grilled sardines on a stick, Málaga beach specialty'),
        ('🥖 Flamenquín', 'Rolled pork filled with ham and cheese, breaded and fried'),
        ('🥪 Pringá', 'Slow-cooked meat sandwich, popular in Seville'),
        ('🍩 Churros con Chocolate', 'Fried dough with thick hot chocolate for dipping'),
        ('🍷 Sherry Wine', 'From Jerez - try fino, manzanilla, or sweet Pedro Ximénez'),
        ('🥣 Ajo Blanco', 'Cold white soup with almonds and grapes, from Málaga')
    ]
    
    for dish_emoji_name, description in dishes:
        dish_para = doc.add_paragraph()
        
        dish_name = dish_para.add_run(dish_emoji_name)
        dish_name.bold = True
        dish_name.font.size = Pt(12)
        dish_name.font.color.rgb = RGBColor(230, 126, 34)
        
        dish_para.add_run(f'\n   {description}')
        
        doc.add_paragraph()
    
    # Food timing tip box
    tip_box = doc.add_paragraph()
    tip_box.paragraph_format.left_indent = Inches(0.5)
    tip_box.paragraph_format.right_indent = Inches(0.5)
    
    tip_title = tip_box.add_run('\n⏰  MEAL TIMING IN SPAIN:\n')
    tip_title.font.bold = True
    tip_title.font.color.rgb = RGBColor(231, 76, 60)
    tip_title.font.size = Pt(12)
    
    tip_text = tip_box.add_run('''
    • Breakfast: 8-10am (coffee & pastry)
    • Lunch: 2-4pm (main meal of the day!)
    • Tapas: 8-10pm (pre-dinner snacks)
    • Dinner: 9pm-midnight (yes, really!)
    
    Pro tip: When in doubt, follow the locals! If a restaurant is full of Spanish families at 10pm, you're in the right place. 🎯
    ''')
    tip_text.font.size = Pt(10)
    tip_text.font.color.rgb = RGBColor(52, 73, 94)
    
    # ========================================================================
    # 🚗 CAR-SPECIFIC TIPS (if road trip)
    # ========================================================================
    
    if is_car_mode:
        doc.add_page_break()
        
        car_header = doc.add_heading('', 1)
        car_run = car_header.add_run('🚗  ROAD TRIP ESSENTIALS')
        car_run.font.size = Pt(24)
        car_run.font.color.rgb = RGBColor(231, 76, 60)
        
        doc.add_paragraph('═' * 50)
        
        # Driving basics
        basics_header = doc.add_heading('', 2)
        basics_run = basics_header.add_run('🛣️  Driving in Spain 101')
        basics_run.font.size = Pt(18)
        basics_run.font.color.rgb = RGBColor(52, 152, 219)
        
        driving_basics = [
            '🚗  Drive on the RIGHT side of the road',
            '🚦  Speed limits: 120 km/h (highways), 90 km/h (rural), 50 km/h (cities)',
            '🚫  Right-of-way: Traffic from right has priority',
            '📱  Using phones while driving is ILLEGAL (€200 fine!)',
            '🍺  Blood alcohol limit: 0.5g/l (0.25g/l for new drivers)',
            '👶  Children under 135cm MUST use car seats',
            '🔺  Required: 2 warning triangles, reflective vest, spare tire',
            '💡  Headlights required in tunnels and at night'
        ]
        
        for tip in driving_basics:
            tip_para = doc.add_paragraph(tip, style='List Bullet')
            tip_para.runs[0].font.size = Pt(10)
            tip_para.runs[0].font.color.rgb = RGBColor(52, 73, 94)
        
        doc.add_paragraph()
        
        # Tolls & Fuel
        tolls_header = doc.add_heading('', 2)
        tolls_run = tolls_header.add_run('💶  Tolls & Fuel')
        tolls_run.font.size = Pt(18)
        tolls_run.font.color.rgb = RGBColor(46, 204, 113)
        
        toll_tips = [
            '🛣️  Autopistas (AP-) are TOLL roads, Autovías (A-) are FREE',
            '💳  Most toll booths accept credit cards (Visa/Mastercard)',
            '🧾  Typical tolls: Málaga↔Seville ~€15-20',
            '⛽  Gas stations (gasolineras) are frequent on highways',
            '💳  Most accept credit cards, some require chip & PIN',
            '⛽  Fuel types: Gasolina 95 (regular), Gasolina 98 (premium), Diesel',
            '📱  Apps: Google Maps or Waze for cheapest fuel nearby'
        ]
        
        for tip in toll_tips:
            tip_para = doc.add_paragraph(tip, style='List Bullet')
            tip_para.runs[0].font.size = Pt(10)
            tip_para.runs[0].font.color.rgb = RGBColor(52, 73, 94)
        
        doc.add_paragraph()
        
        # Parking
        parking_header = doc.add_heading('', 2)
        parking_run = parking_header.add_run('🅿️  Parking Guide')
        parking_run.font.size = Pt(18)
        parking_run.font.color.rgb = RGBColor(155, 89, 182)
        
        parking_intro = doc.add_paragraph()
        parking_intro_run = parking_intro.add_run('Parking in historic centers can be tricky! Here\'s what the colors mean:')
        parking_intro_run.font.size = Pt(10)
        parking_intro_run.italic = True
        parking_intro_run.font.color.rgb = RGBColor(127, 140, 141)
        
        doc.add_paragraph()
        
        parking_tips = [
            '🔵  BLUE zones: Pay & display (limited time, usually 2-4h)',
            '🟢  GREEN zones: Resident parking ONLY - avoid these!',
            '⚪  WHITE lines: Free parking (rare in city centers)',
            '🟡  YELLOW lines: No parking/stopping zones',
            '🅿️  Public lots: €15-25/day in city centers',
            '🏨  Hotel parking: €10-20/night (ask when booking)',
            '💡  Best strategy: Park outside old town, walk or taxi in',
            '📱  Useful apps: Parclick, ElParking, Parkopedia',
            '⚠️  NEVER leave valuables visible in car!'
        ]
        
        for tip in parking_tips:
            tip_para = doc.add_paragraph(tip, style='List Bullet')
            tip_para.runs[0].font.size = Pt(10)
            tip_para.runs[0].font.color.rgb = RGBColor(52, 73, 94)
        
        doc.add_paragraph()
        
        # Car rental tips
        rental_header = doc.add_heading('', 2)
        rental_run = rental_header.add_run('🔑  Car Rental Insider Tips')
        rental_run.font.size = Pt(18)
        rental_run.font.color.rgb = RGBColor(243, 156, 18)
        
        rental_tips = [
            '📅  Book online in advance = 30-50% cheaper!',
            '🔞  Most rentals require 21+ (sometimes 25+)',
            '💳  Credit card required for deposit (debit often not accepted)',
            '🛡️  Consider full insurance (CDW + theft protection)',
            '📄  Bring: Valid license, passport, credit card',
            '🌍  International Driving Permit: Recommended for non-EU licenses',
            '⛽  "Full to full" policy is standard (return with full tank)',
            '🚗  Manual transmission is default - automatic costs MORE',
            '📱  GPS: €5-10/day, or just use your phone',
            '📸  Photograph ANY existing damage before leaving lot!'
        ]
        
        for tip in rental_tips:
            tip_para = doc.add_paragraph(tip, style='List Bullet')
            tip_para.runs[0].font.size = Pt(10)
            tip_para.runs[0].font.color.rgb = RGBColor(52, 73, 94)
    
    # ========================================================================
    # 💡 GENERAL TRAVEL TIPS
    # ========================================================================
    
    doc.add_page_break()
    
    general_header = doc.add_heading('', 1)
    general_run = general_header.add_run('💡  ESSENTIAL TRAVEL TIPS')
    general_run.font.size = Pt(24)
    general_run.font.color.rgb = RGBColor(52, 152, 219)
    
    doc.add_paragraph('═' * 50)
    
    general_tips = [
        ('⏰ Spanish Schedule', 'Lunch 2-4pm, Dinner 9pm+. Many restaurants closed 4-8pm. Embrace the siesta!'),
        ('🎫 Book Ahead', 'Alhambra, Alcázar, Cathedral tours need 2-3 weeks advance booking!'),
        ('☀️ Best Season', 'Spring (April-May) or Fall (September-October) for perfect weather'),
        ('💶 Cash is King', 'Small towns, markets, and tapas bars often prefer cash'),
        ('🗣️ Learn Spanish', 'Even basic phrases go a long way! Locals really appreciate it'),
        ('🍷 Free Tapas', 'In Granada, many bars give FREE tapas with drinks!'),
        ('🏛️ Monday Closures', 'Many museums closed Mondays - plan accordingly'),
        ('📱 Get Data', 'Local SIM or international plan for navigation & translations'),
        ('👟 Comfy Shoes', 'Cobblestone streets everywhere - sneakers are your friend'),
        ('🌡️ Summer Heat', 'June-August = 40°C+. Plan indoor activities for midday'),
        ('🚰 Tap Water', 'Safe to drink, but locals prefer bottled'),
        ('⏰ Siesta Time', 'Small shops close 2-5pm (tourist areas stay open)')
    ]
    
    for tip_title, tip_desc in general_tips:
        tip_para = doc.add_paragraph()
        
        title_run = tip_para.add_run(tip_title)
        title_run.bold = True
        title_run.font.size = Pt(11)
        title_run.font.color.rgb = RGBColor(41, 128, 185)
        
        desc_run = tip_para.add_run(f'\n   {tip_desc}')
        desc_run.font.size = Pt(10)
        desc_run.font.color.rgb = RGBColor(52, 73, 94)
        
        doc.add_paragraph()
    
    # ========================================================================
    # 🎒 PACKING LIST
    # ========================================================================
    
    doc.add_page_break()
    
    packing_header = doc.add_heading('', 1)
    packing_run = packing_header.add_run('🎒  PACKING CHECKLIST')
    packing_run.font.size = Pt(24)
    packing_run.font.color.rgb = RGBColor(155, 89, 182)
    
    doc.add_paragraph('═' * 50)
    
    packing_intro = doc.add_paragraph()
    packing_intro_run = packing_intro.add_run('Pack smart for your Andalusian adventure! Here\'s everything you need:')
    packing_intro_run.font.size = Pt(11)
    packing_intro_run.italic = True
    packing_intro_run.font.color.rgb = RGBColor(127, 140, 141)
    
    doc.add_paragraph()
    
    packing_categories = {
        '👕 Clothing': [
            '👟  Comfortable walking shoes (10-20km per day!)',
            '🩴  Sandals or flip-flops for beach/hotel',
            '👕  Light, breathable clothing (cotton/linen)',
            '🧥  Light jacket for evenings (even in summer)',
            '👗  Modest clothes for churches (covered shoulders/knees)',
            '🩱  Swimsuit for beaches or hotel pools',
            '🧢  Hat or cap for sun protection'
        ],
        '🎒 Essentials': [
            '🧴  Sunscreen SPF 30+ (Andalusian sun is STRONG!)',
            '🕶️  Sunglasses with UV protection',
            '💧  Reusable water bottle',
            '🎒  Day backpack for sightseeing',
            '🔌  Power adapter (Type C/F for Spain)',
            '🔋  Portable phone charger',
            '📄  Copy of passport & travel insurance',
            '💊  Prescription meds + basic first aid'
        ],
        '✨ Nice to Have': [
            '🌧️  Light rain jacket (spring/fall)',
            '📖  Spanish phrasebook or translation app',
            '📷  Camera with good storage',
            '👔  Smart-casual outfit for nice dinners',
            '🔒  Small lock for hotel lockers',
            '👂  Earplugs (Spanish cities = noisy at night!)',
            '🧼  Hand sanitizer and wet wipes'
        ]
    }
    
    for category, items in packing_categories.items():
        cat_header = doc.add_heading('', 2)
        cat_run = cat_header.add_run(category)
        cat_run.font.size = Pt(16)
        cat_run.font.color.rgb = RGBColor(52, 152, 219)
        
        for item in items:
            item_para = doc.add_paragraph(item, style='List Bullet')
            item_para.runs[0].font.size = Pt(10)
            item_para.runs[0].font.color.rgb = RGBColor(52, 73, 94)
        
        doc.add_paragraph()
    
    # ========================================================================
    # 🗣️ USEFUL SPANISH PHRASES
    # ========================================================================
    
    doc.add_page_break()
    
    phrases_header = doc.add_heading('', 1)
    phrases_run = phrases_header.add_run('🗣️  SURVIVAL SPANISH')
    phrases_run.font.size = Pt(24)
    phrases_run.font.color.rgb = RGBColor(230, 126, 34)
    
    doc.add_paragraph('═' * 50)
    
    phrases_intro = doc.add_paragraph()
    phrases_intro_run = phrases_intro.add_run('Basic Spanish will make your trip SO much better! Practice these:')
    phrases_intro_run.font.size = Pt(11)
    phrases_intro_run.italic = True
    phrases_intro_run.font.color.rgb = RGBColor(127, 140, 141)
    
    doc.add_paragraph()
    
    phrases = [
        ('👋  Hola / Buenos días', 'Hello / Good morning'),
        ('🙏  Gracias / Muchas gracias', 'Thank you / Thank you very much'),
        ('🙂  Por favor', 'Please'),
        ('😊  De nada', 'You\'re welcome'),
        ('❓  ¿Habla inglés?', 'Do you speak English?'),
        ('🤷  No entiendo', 'I don\'t understand'),
        ('💰  ¿Cuánto cuesta?', 'How much does it cost?'),
        ('🧾  La cuenta, por favor', 'The bill, please'),
        ('📍  ¿Dónde está...?', 'Where is...?'),
        ('🍽️  Una mesa para dos', 'A table for two'),
        ('🍻  ¡Salud!', 'Cheers!'),
        ('😅  Perdón / Disculpe', 'Excuse me / Sorry'),
        ('✅  Sí / ❌  No', 'Yes / No'),
        ('👋  Adiós / Hasta luego', 'Goodbye / See you later'),
        ('🚨  ¡Ayuda!', 'Help!'),
        ('🏥  Necesito un médico', 'I need a doctor')
    ]
    
    for emoji_spanish, english in phrases:
        phrase_para = doc.add_paragraph()
        
        spanish_run = phrase_para.add_run(emoji_spanish)
        spanish_run.bold = True
        spanish_run.font.size = Pt(12)
        spanish_run.font.color.rgb = RGBColor(231, 76, 60)
        
        english_run = phrase_para.add_run(f'  →  {english}')
        english_run.font.size = Pt(11)
        english_run.font.color.rgb = RGBColor(52, 73, 94)
    
    # ========================================================================
    # 📞 EMERGENCY CONTACTS
    # ========================================================================
    
    doc.add_page_break()
    
    emergency_header = doc.add_heading('', 1)
    emergency_run = emergency_header.add_run('📞  EMERGENCY CONTACTS')
    emergency_run.font.size = Pt(24)
    emergency_run.font.color.rgb = RGBColor(231, 76, 60)
    
    doc.add_paragraph('═' * 50)
    
    emergency_intro = doc.add_paragraph()
    emergency_intro_run = emergency_intro.add_run('⚠️  Save these numbers in your phone BEFORE you travel!')
    emergency_intro_run.font.bold = True
    emergency_intro_run.font.size = Pt(12)
    emergency_intro_run.font.color.rgb = RGBColor(231, 76, 60)
    
    doc.add_paragraph()
    
    contacts = [
        ('🚨  ALL EMERGENCIES (EU-wide)', '112', 'Police, medical, fire - works everywhere in Europe'),
        ('👮  National Police', '091', 'For crimes, theft, lost documents'),
        ('🚑  Medical Emergency', '061', 'Ambulance and medical assistance'),
        ('🔥  Fire Department', '080', 'Fire emergencies'),
        ('🚓  Local Police', '092', 'Non-emergency local police'),
        ('ℹ️  Tourist Information', '902 200 120', 'Tourism information hotline'),
    ]
    
    for emoji_service, number, description in contacts:
        contact_para = doc.add_paragraph()
        
        service_run = contact_para.add_run(f'{emoji_service}:  ')
        service_run.font.bold = True
        service_run.font.size = Pt(12)
        service_run.font.color.rgb = RGBColor(52, 73, 94)
        
        number_run = contact_para.add_run(number)
        number_run.font.size = Pt(14)
        number_run.font.bold = True
        number_run.font.color.rgb = RGBColor(231, 76, 60)
        
        desc_para = doc.add_paragraph()
        desc_para.paragraph_format.left_indent = Inches(0.5)
        desc_run = desc_para.add_run(f'   {description}')
        desc_run.font.size = Pt(9)
        desc_run.font.color.rgb = RGBColor(127, 140, 141)
        desc_run.italic = True
    
    doc.add_paragraph()
    
    embassy_header = doc.add_heading('', 2)
    embassy_run = embassy_header.add_run('🏛️  Embassy Contacts (Madrid)')
    embassy_run.font.size = Pt(16)
    embassy_run.font.color.rgb = RGBColor(52, 152, 219)
    
    embassies = [
        ('🇺🇸  US Embassy', '+34 91 587 2200'),
        ('🇬🇧  UK Embassy', '+34 91 714 6300'),
        ('🇨🇦  Canadian Embassy', '+34 91 382 8400'),
        ('🇦🇺  Australian Embassy', '+34 91 353 6600'),
        ('🇮🇪  Irish Embassy', '+34 91 436 4093')
    ]
    
    for embassy, phone in embassies:
        emb_para = doc.add_paragraph()
        emb_para.paragraph_format.left_indent = Inches(0.3)
        
        emb_name = emb_para.add_run(f'{embassy}:  ')
        emb_name.font.bold = True
        emb_name.font.size = Pt(11)
        emb_name.font.color.rgb = RGBColor(52, 73, 94)
        
        emb_phone = emb_para.add_run(phone)
        emb_phone.font.size = Pt(11)
        emb_phone.font.color.rgb = RGBColor(41, 128, 185)
    
    # ========================================================================
    # ✨ BEAUTIFUL CLOSING PAGE
    # ========================================================================
    
    doc.add_page_break()
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Final inspiring message
    closing_header = doc.add_paragraph()
    closing_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    closing_run = closing_header.add_run('✨  HAVE AN AMAZING ADVENTURE!  ✨')
    closing_run.font.size = Pt(28)
    closing_run.font.color.rgb = RGBColor(41, 128, 185)
    closing_run.bold = True
    
    doc.add_paragraph()
    
    # Closing quote
    closing_quote = doc.add_paragraph()
    closing_quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
    quote_run = closing_quote.add_run('"Travel is the only thing you buy that makes you richer."')
    quote_run.font.size = Pt(14)
    quote_run.italic = True
    quote_run.font.color.rgb = RGBColor(127, 140, 141)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Andalusia flag emojis
    flag_para = doc.add_paragraph()
    flag_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    flag_run = flag_para.add_run('🌊 ☀️ 🏰 🍷 🎭 🎸 💃')
    flag_run.font.size = Pt(24)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Enjoy message
    enjoy_para = doc.add_paragraph()
    enjoy_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    enjoy_run = enjoy_para.add_run('Enjoy every moment in beautiful Andalusia!\nMake memories, take photos, eat tapas, and embrace the adventure.')
    enjoy_run.font.size = Pt(12)
    enjoy_run.font.color.rgb = RGBColor(52, 73, 94)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Hashtag fun
    hashtag_para = doc.add_paragraph()
    hashtag_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hashtag_run = hashtag_para.add_run('#AndalusiaRoadTrip #TravelSpain #Wanderlust')
    hashtag_run.font.size = Pt(10)
    hashtag_run.font.color.rgb = RGBColor(149, 165, 166)
    hashtag_run.italic = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # ✅ NEW: Affiliate disclosure (required by law in most jurisdictions)
    if BOOKING_AFFILIATE_ID and BOOKING_AFFILIATE_ID != "YOUR_AFFILIATE_ID":
        disclosure = doc.add_paragraph()
        disclosure.alignment = WD_ALIGN_PARAGRAPH.CENTER
        disclosure_run = disclosure.add_run(
            '📝 Disclosure: This itinerary contains affiliate links to booking partners. '
            'If you book through these links, we may earn a commission at no extra cost to you. '
            'This helps us keep our service free and improve our recommendations. Thank you for your support!'
        )
        disclosure_run.font.size = Pt(7)
        disclosure_run.italic = True
        disclosure_run.font.color.rgb = RGBColor(149, 165, 166)
        
        doc.add_paragraph()
    
    # ========================================================================
    # 🎉 EVENTS SECTION
    # ========================================================================
    
    # Get events for the trip
    events_list = []
    
    if start_date:
        try:
            # Import events service
            from events_service import get_events_for_trip
            from datetime import timedelta
            
            # Get end date
            end_date = start_date + timedelta(days=days - 1)
            
            # ✅ UPDATED: Check ALL major Andalusia cities (same as UI)
            major_cities = ['Seville', 'Granada', 'Córdoba', 'Málaga', 'Cádiz', 'Jerez', 'Ronda']
            cities_to_check = list(set([c.title() for c in ordered_cities] + major_cities))
            
            # Fetch events for each city
            for city in cities_to_check:
                try:
                    city_events = get_events_for_trip(
                        city,
                        start_date.strftime('%Y-%m-%d'),
                        end_date.strftime('%Y-%m-%d'),
                        eventbrite_token=None
                    )
                    events_list.extend(city_events)
                except Exception as e:
                    # print(f"⚠️ Error fetching events for {city} in document: {e}")
                    continue
            
            # Remove duplicates by name+date
            seen = set()
            unique_events = []
            for event in events_list:
                key = (event.get('name'), event.get('date'))
                if key not in seen:
                    seen.add(key)
                    unique_events.append(event)
            events_list = unique_events
            
            # Sort by tier (tier_1 first) then by date
            tier_order = {'tier_1': 0, 'tier_2': 1, 'tier_3': 2}
            events_list.sort(key=lambda x: (tier_order.get(x.get('tier', 'tier_3'), 3), x.get('date', '')))
            
        except Exception as e:
            # print(f"⚠️ Could not fetch events for document: {e}")
            events_list = []
    
    # Display events if found
    if events_list:
        doc.add_page_break()
        
        # Events header
        events_header = doc.add_heading('', 1)
        events_run = events_header.add_run('🎉  EVENTS DURING YOUR TRIP')
        events_run.font.size = Pt(24)
        events_run.font.color.rgb = RGBColor(230, 126, 34)
        events_run.bold = True
        
        separator = doc.add_paragraph('─' * 50)
        separator_run = separator.runs[0]
        separator_run.font.color.rgb = RGBColor(189, 195, 199)
        
        # Intro text
        intro = doc.add_paragraph()
        intro_run = intro.add_run(
            'Special events and festivals happening during your visit. '
            'Plan ahead to experience local culture at its best!'
        )
        intro_run.font.size = Pt(11)
        intro_run.italic = True
        intro_run.font.color.rgb = RGBColor(127, 140, 141)
        
        doc.add_paragraph()
        
        # Display each event
        for idx, event in enumerate(events_list[:10], 1):  # Show up to 10 events
            event_date = event.get('date', '')
            event_name = event.get('name', 'Event')
            event_location = event.get('location', '')
            event_desc = event.get('description', '')
            event_type = event.get('type', '')
            event_source = event.get('source', '')
            
            # ============================================================
            # 📸 TRY TO ADD EVENT PHOTO
            # ============================================================
            event_photo_path = None
            
            # DEBUG: Show current working directory
            import os
            
            # Try to find photo based on event name
            # Look in event_photos directory using portable path
            if event_name:
                # Normalize event name for filename
                event_filename = event_name.lower().replace(' ', '_').replace('(', '').replace(')', '')
                event_filename = event_filename.replace('á', 'a').replace('é', 'e').replace('í', 'i').replace('ó', 'o').replace('ú', 'u')
                
                # Try multiple possible locations using EVENT_PHOTOS_DIR
                possible_paths = [
                    os.path.join(EVENT_PHOTOS_DIR, f"{event_filename}.jpg"),
                    os.path.join(EVENT_PHOTOS_DIR, f"{event_filename}.png"),
                    os.path.join(EVENT_PHOTOS_DIR, f"{event_type.lower()}.jpg"),
                    os.path.join(EVENT_PHOTOS_DIR, f"{event_type.lower()}.png"),
                    os.path.join(EVENT_PHOTOS_DIR, f"{event_location.lower()}.jpg"),
                    os.path.join(EVENT_PHOTOS_DIR, f"{event_location.lower()}.png"),
                ]
                
                for path in possible_paths:
                    exists = os.path.exists(path)
                    if exists:
                        event_photo_path = path
                        break
            
            # If photo found, add it
            if event_photo_path:
                try:
                    photo_para = doc.add_paragraph()
                    photo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    photo_para.paragraph_format.space_before = Pt(12)
                    photo_para.paragraph_format.space_after = Pt(6)
                    
                    photo_run = photo_para.add_run()
                    photo_run.add_picture(event_photo_path, width=Inches(5.0))
                    
                    # Photo caption
                    caption_para = doc.add_paragraph()
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_run = caption_para.add_run(event_name)
                    caption_run.font.size = Pt(9)
                    caption_run.italic = True
                    caption_run.font.color.rgb = RGBColor(127, 140, 141)
                    
                    doc.add_paragraph()  # Space after photo
                except Exception as e:
                    print(f"⚠️ Could not add event photo: {e}")
            
            # ============================================================
            # EVENT TEXT CONTENT
            # ============================================================
            
            # Event container with subtle background (using table for better styling)
            event_para = doc.add_paragraph()
            event_para.paragraph_format.left_indent = Inches(0.3)
            event_para.paragraph_format.space_before = Pt(12)
            event_para.paragraph_format.space_after = Pt(12)
            
            # Icon based on type (large and colorful)
            icon_map = {
                'Festival': '🎊',
                'Concert': '🎵',
                'Sports': '⚽',
                'Music': '🎭',
                'Religious': '⛪',
                'Cultural': '🎨',
                'Flamenco': '💃',
                'Carnival': '🎭',
                'Theater': '🎭'
            }
            icon = icon_map.get(event_type, '🎉')
            
            # Event title with large icon
            event_title_run = event_para.add_run(f'{icon}  {event_name}')
            event_title_run.font.size = Pt(14)
            event_title_run.bold = True
            event_title_run.font.color.rgb = RGBColor(52, 73, 94)
            
            # Event details in a box
            details_para = doc.add_paragraph()
            details_para.paragraph_format.left_indent = Inches(0.5)
            details_para.paragraph_format.space_before = Pt(3)
            
            # Date and location with colored background effect (using symbols)
            details_run = details_para.add_run(f'📅 {event_date}  •  📍 {event_location}  •  🏷️ {event_type}')
            details_run.font.size = Pt(10)
            details_run.font.color.rgb = RGBColor(149, 165, 166)
            
            # Description
            if event_desc:
                desc_para = doc.add_paragraph()
                desc_para.paragraph_format.left_indent = Inches(0.5)
                desc_para.paragraph_format.space_before = Pt(6)
                
                # Truncate if too long
                if len(event_desc) > 250:
                    event_desc = event_desc[:250] + '...'
                
                desc_run = desc_para.add_run(event_desc)
                desc_run.font.size = Pt(10)
                desc_run.font.color.rgb = RGBColor(52, 73, 94)
            
            # Add visual separator between description and source
            separator_para = doc.add_paragraph()
            separator_para.paragraph_format.left_indent = Inches(0.5)
            separator_para.paragraph_format.space_before = Pt(6)
            separator_run = separator_para.add_run('─' * 60)
            separator_run.font.size = Pt(8)
            separator_run.font.color.rgb = RGBColor(220, 220, 220)
            
            # Source attribution with icon
            if event_source:
                source_para = doc.add_paragraph()
                source_para.paragraph_format.left_indent = Inches(0.5)
                source_para.paragraph_format.space_before = Pt(3)
                source_run = source_para.add_run(f'ℹ️ Source: {event_source}')
                source_run.font.size = Pt(8)
                source_run.italic = True
                source_run.font.color.rgb = RGBColor(189, 195, 199)
            
            # Tip for major festivals
            if event_type in ['Festival', 'Carnival'] or 'Feria' in event_name or 'Festival' in event_name:
                tip_para = doc.add_paragraph()
                tip_para.paragraph_format.left_indent = Inches(0.5)
                tip_para.paragraph_format.space_before = Pt(6)
                tip_run = tip_para.add_run('💡 Tip: Book accommodations early - prices increase during festivals!')
                tip_run.font.size = Pt(9)
                tip_run.italic = True
                tip_run.font.color.rgb = RGBColor(241, 196, 15)  # Golden yellow
            
            doc.add_paragraph()  # Space between events
        
        # Show count if more events exist
        if len(events_list) > 10:
            more_para = doc.add_paragraph()
            more_run = more_para.add_run(f'... and {len(events_list) - 10} more events!')
            more_run.font.size = Pt(10)
            more_run.italic = True
            more_run.font.color.rgb = RGBColor(127, 140, 141)
        
        doc.add_paragraph()
    
    # Generated by footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run('Generated with ❤️ by Your Personal Travel Planner')
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(189, 195, 199)
    
    # Save to BytesIO
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio